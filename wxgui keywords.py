import wx
import os
from indicnlp.tokenize import indic_tokenize
from collections import Counter
from docx import Document
from wx import adv
from fpdf import FPDF
import shutil
import fitz  # PyMuPDF
from stopwords_kannada.stopwords import stopword
import pytesseract



class UnicodePDF(FPDF):
    def header(self):
        pass

    def footer(self):
        pass


class KannadaKeywordExtractorApp(wx.Frame):
    def __init__(self, parent, title):
        super(KannadaKeywordExtractorApp, self).__init__(parent, title=title, size=(700, 500))

        # Initialize output file paths
        self.word_output_file = None
        self.pdf_output_file = None

        # Initialize category output files
        self.noun_output_file = None
        self.verb_output_file = None
        self.pronoun_output_file = None

        # Default number of keywords to extract
        self.num_keywords = "5"  # Default to extracting 5 keywords

        # Create and set up the GUI elements
        self.init_ui()

    def init_ui(self):
        panel = wx.Panel(self)

        # File selection area
        file_label = wx.StaticText(panel, label="Select Document:")
        self.file_entry = wx.TextCtrl(panel, style=wx.TE_READONLY)
        browse_button = wx.Button(panel, label="Browse")
        browse_button.Bind(wx.EVT_BUTTON, self.on_browse)

        # Text input area
        text_label = wx.StaticText(panel, label="Enter Kannada Text:")
        self.text_input = wx.TextCtrl(panel, style=wx.TE_MULTILINE)

        # Number of keywords entry
        num_keywords_label = wx.StaticText(panel, label="Number of Keywords:")
        self.num_keywords_entry = wx.TextCtrl(panel, value=self.num_keywords)

        # Button to trigger keyword extraction
        extract_button = wx.Button(panel, label="Extract Keywords")
        extract_button.Bind(wx.EVT_BUTTON, self.extract_keywords)

        # Display area for extracted keywords
        output_label = wx.StaticText(panel, label="Extracted Keywords:")
        self.keywords_display = wx.TextCtrl(panel, style=wx.TE_MULTILINE | wx.TE_READONLY,
                                            size=(200, 100))  # Adjust size here

        # Display area for extracted nouns, pronouns, and verbs
        noun_label = wx.StaticText(panel, label="ನಾಮಪದ:")
        self.noun_display = wx.TextCtrl(panel, style=wx.TE_MULTILINE | wx.TE_READONLY, size=(200, 100))  # Larger size
        verb_label = wx.StaticText(panel, label="ಕ್ರಿಯಾಪದ:")
        self.verb_display = wx.TextCtrl(panel, style=wx.TE_MULTILINE | wx.TE_READONLY, size=(200, 100))  # Larger size
        pronoun_label = wx.StaticText(panel, label="ಸರ್ವನಾಮಗಳು:")
        self.pronoun_display = wx.TextCtrl(panel, style=wx.TE_MULTILINE | wx.TE_READONLY,
                                           size=(200, 100))  # Larger size

        # Buttons for downloading files
        download_word_button = wx.Button(panel, label="Download Word")
        download_word_button.Bind(wx.EVT_BUTTON, self.download_word)
        download_pdf_button = wx.Button(panel, label="Download PDF")
        download_pdf_button.Bind(wx.EVT_BUTTON, self.download_pdf)
        download_all_button = wx.Button(panel, label="Download All Categories")
        download_all_button.Bind(wx.EVT_BUTTON, self.download_all_categories)

        # Clear button
        clear_button = wx.Button(panel, label="Clear")
        clear_button.Bind(wx.EVT_BUTTON, self.clear_all)

        # Separate action buttons for noun, verb, and pronoun extraction
        extract_noun_button = wx.Button(panel, label="Extract ನಾಮಪದ")
        extract_noun_button.Bind(wx.EVT_BUTTON, self.extract_nouns)

        extract_verb_button = wx.Button(panel, label="Extract ಕ್ರಿಯಾಪದ")
        extract_verb_button.Bind(wx.EVT_BUTTON, self.extract_verbs)

        extract_pronoun_button = wx.Button(panel, label="Extract ಸರ್ವನಾಮಗಳು")
        extract_pronoun_button.Bind(wx.EVT_BUTTON, self.extract_pronouns)

        # Sizers for layout
        sizer = wx.GridBagSizer(5, 5)
        sizer.Add(file_label, (0, 0))
        sizer.Add(self.file_entry, (0, 1), (1, 2), flag=wx.EXPAND)
        sizer.Add(browse_button, (0, 3))
        sizer.Add(text_label, (1, 0))
        sizer.Add(self.text_input, (1, 1), (2, 3), flag=wx.EXPAND)
        sizer.Add(num_keywords_label, (3, 0))
        sizer.Add(self.num_keywords_entry, (3, 1))
        sizer.Add(extract_button, (3, 2))
        sizer.Add(output_label, (4, 0))
        sizer.Add(self.keywords_display, (5, 0), (2, 4), flag=wx.EXPAND)
        sizer.Add(noun_label, (7, 0))
        sizer.Add(self.noun_display, (8, 0), (2, 1), flag=wx.EXPAND)
        sizer.Add(verb_label, (7, 1))
        sizer.Add(self.verb_display, (8, 1), (2, 1), flag=wx.EXPAND)
        sizer.Add(pronoun_label, (7, 2))
        sizer.Add(self.pronoun_display, (8, 2), (2, 1), flag=wx.EXPAND)
        sizer.Add(download_word_button, (10, 0))
        sizer.Add(download_pdf_button, (10, 1))
        sizer.Add(download_all_button, (10, 2))
        sizer.Add(clear_button, (10, 3))

        # Add action buttons to the sizer
        sizer.Add(extract_noun_button, (11, 0))
        sizer.Add(extract_verb_button, (11, 1))
        sizer.Add(extract_pronoun_button, (11, 2))

        panel.SetSizer(sizer)

        self.Centre()
        self.Show(True)

    def on_browse(self, event):
        wildcard = "All files (*.*)|*.*"
        dialog = wx.FileDialog(self, "Choose a file", style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST, wildcard=wildcard)
        if dialog.ShowModal() == wx.ID_OK:
            self.file_entry.SetValue(dialog.GetPath())
        dialog.Destroy()

    def extract_keywords(self, event):
        # Get document file path
        doc_file_path = self.file_entry.GetValue()

        # Check if the text is entered manually
        if not doc_file_path:
            doc_text = self.text_input.GetValue()
        else:
            # Extract text from document
            doc_text = self.extract_text_from_document(doc_file_path)

            # Display extracted text in the text input area
            self.text_input.SetValue(doc_text)

        # Extract keywords
        num_keywords = int(self.num_keywords_entry.GetValue())
        noun_words, verb_words, pronoun_words = self.extract_kannada_keywords(doc_text, num_keywords)

        # Display extracted keywords only in the "Extracted Keywords" box
        all_keywords = noun_words + verb_words + pronoun_words
        keywords_str = ', '.join(all_keywords)
        self.keywords_display.SetValue(keywords_str)

        # Save keywords to output Word and PDF files
        self.save_keywords_to_word(doc_file_path, all_keywords)
        self.save_keywords_to_pdf(doc_file_path, all_keywords)

    def clear_all(self, event):
        # Clear input text area, keywords display area, and selected file entry
        self.text_input.SetValue("")
        self.keywords_display.SetValue("")
        self.file_entry.SetValue("")
        self.noun_display.SetValue("")
        self.verb_display.SetValue("")
        self.pronoun_display.SetValue("")

    def extract_nouns(self, event):
        # Get document file path
        doc_file_path = self.file_entry.GetValue()

        # Check if the text is entered manually
        if not doc_file_path:
            doc_text = self.text_input.GetValue()
        else:
            # Extract text from document
            doc_text = self.extract_text_from_document(doc_file_path)

            # Display extracted text in the text input area
            self.text_input.SetValue(doc_text)

        # Extract keywords
        num_keywords = int(self.num_keywords_entry.GetValue())
        noun_words, _, _ = self.extract_kannada_keywords(doc_text, num_keywords)

        # Display extracted keywords in the "Nouns" box
        self.noun_display.SetValue(', '.join(noun_words))

    def extract_verbs(self, event):
        # Get document file path
        doc_file_path = self.file_entry.GetValue()

        # Check if the text is entered manually
        if not doc_file_path:
            doc_text = self.text_input.GetValue()
        else:
            # Extract text from document
            doc_text = self.extract_text_from_document(doc_file_path)

            # Display extracted text in the text input area
            self.text_input.SetValue(doc_text)

        # Extract keywords
        num_keywords = int(self.num_keywords_entry.GetValue())
        _, verb_words, _ = self.extract_kannada_keywords(doc_text, num_keywords)

        # Display extracted keywords in the "Verbs" box
        self.verb_display.SetValue(', '.join(verb_words))

    def extract_pronouns(self, event):
        # Get document file path
        doc_file_path = self.file_entry.GetValue()

        # Check if the text is entered manually
        if not doc_file_path:
            doc_text = self.text_input.GetValue()
        else:
            # Extract text from document
            doc_text = self.extract_text_from_document(doc_file_path)

            # Display extracted text in the text input area
            self.text_input.SetValue(doc_text)

        # Extract keywords
        num_keywords = int(self.num_keywords_entry.GetValue())
        _, _, pronoun_words = self.extract_kannada_keywords(doc_text, num_keywords)

        # Display extracted keywords in the "Pronouns" box
        self.pronoun_display.SetValue(', '.join(pronoun_words))

    def extract_text_from_pdf(self, pdf_file_path):
        text = ""
        try:
            doc = fitz.open(pdf_file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
        return text

    def extract_text_from_document(self, doc_file_path):
        if doc_file_path.endswith('.txt'):
            with open(doc_file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif doc_file_path.endswith('.docx'):
            doc = Document(doc_file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        elif doc_file_path.endswith('.pdf'):
            return self.extract_text_from_pdf(doc_file_path)
        else:
            return ''

    def save_keywords_to_word(self, doc_file_path, keywords):
        if doc_file_path.endswith('.txt') or doc_file_path.endswith('.docx'):
            output_file_path = doc_file_path.replace('.txt', '_keywords.docx').replace('.docx', '_keywords.docx')
            doc = Document()
            doc.add_heading('Extracted Keywords', level=1)

            # Add extracted keywords to the document
            for keyword in keywords:
                doc.add_paragraph(keyword)

            doc.save(output_file_path)
            self.word_output_file = output_file_path

    def save_keywords_to_pdf(self, doc_file_path, keywords):
        if doc_file_path.endswith('.txt') or doc_file_path.endswith('.docx'):
            output_file_path = doc_file_path.replace('.txt', '_keywords.pdf').replace('.docx', '_keywords.pdf')
            pdf = UnicodePDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            # Add extracted keywords to the PDF
            for keyword in keywords:
                try:
                    pdf.cell(200, 10, txt=keyword, ln=True, align='L')
                except UnicodeEncodeError as e:
                    # Handle encoding issues, you can log the error or handle it as needed
                    print(f"UnicodeEncodeError: {e}")

            try:
                with open(output_file_path, 'wb') as pdf_file:
                    pdf_file.write(pdf.output(dest='S').encode('utf-8'))  # Changed encoding to UTF-8
                self.pdf_output_file = output_file_path
            except Exception as e:
                print(f"Error saving PDF file: {e}")

    def save_category_to_word(self, category_name, keywords):
        if category_name:
            output_file_path = f"{category_name}_keywords.docx"
            doc = Document()
            doc.add_heading(f'Extracted {category_name}', level=1)

            # Add extracted keywords to the document
            for keyword in keywords:
                doc.add_paragraph(keyword)

            doc.save(output_file_path)
            if category_name == "Nouns":
                self.noun_output_file = output_file_path
            elif category_name == "Verbs":
                self.verb_output_file = output_file_path
            elif category_name == "Pronouns":
                self.pronoun_output_file = output_file_path

    def download_word(self, event):
        try:
            if self.word_output_file:
                dialog = wx.FileDialog(self, message="Save Word file", wildcard="Word files (*.docx)|*.docx",
                                       style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT)
                if dialog.ShowModal() == wx.ID_OK:
                    shutil.copyfile(self.word_output_file, dialog.GetPath())
                    self.word_output_file = None  # Reset output file path
        except Exception as e:
            print(f"Error downloading Word file: {e}")

    def download_pdf(self, event):
        try:
            if self.pdf_output_file:
                dialog = adv.FileDialog(self, "Save PDF file", wildcard="PDF files (*.pdf)|*.pdf",
                                        style=adv.FD_SAVE | adv.FD_OVERWRITE_PROMPT)
                if dialog.ShowModal() == wx.ID_CANCEL:
                    dialog.Destroy()
                    return
                else:
                    output_path = dialog.GetPath()
                    shutil.copyfile(self.pdf_output_file, output_path)
                    self.pdf_output_file = None  # Reset output file path
                    dialog.Destroy()
        except Exception as e:
            print(f"Error downloading PDF file: {e}")

    def download_all_categories(self, event):
        try:
            output_file_path = "All_Categories_Keywords.docx"
            doc = Document()

            # Add headings for each category
            if self.noun_output_file and os.path.exists(self.noun_output_file):
                doc.add_heading('Nouns', level=1)
                self.add_content_to_document(doc, self.noun_output_file)

            if self.verb_output_file and os.path.exists(self.verb_output_file):
                doc.add_heading('Verbs', level=1)
                self.add_content_to_document(doc, self.verb_output_file)

            if self.pronoun_output_file and os.path.exists(self.pronoun_output_file):
                doc.add_heading('Pronouns', level=1)
                self.add_content_to_document(doc, self.pronoun_output_file)

            doc.save(output_file_path)
        except Exception as e:
            print(f"Error downloading all categories file: {e}")

    def add_content_to_document(self, doc, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()
            for line in lines:
                doc.add_paragraph(line.strip())

    def extract_kannada_keywords(self, text, num_keywords):
        # Tokenize Kannada text
        tokens = indic_tokenize.trivial_tokenize(text, lang='kn')

        # Filter out stopwords (common words that don't carry much meaning)
        stopwords = stopword()
        filtered_tokens = [token for token in tokens if token not in stopwords]

        # Calculate word frequencies
        word_frequencies = Counter(filtered_tokens)

        # Extract top 'num_keywords' keywords based on frequency
        keywords = [word for word, _ in word_frequencies.most_common(num_keywords)]

        # Separate keywords into noun, verb, and pronoun categories
        noun_words = []
        verb_words = []
        pronoun_words = []

        # for keyword in keywords:
        #     # Assume all keywords are nouns by default
        #     keyword_type = None
        #
        #     # Check if the keyword is a noun based on specific conditions
        #     if keyword.endswith('ವರು') or keyword.endswith('ವಳು') or keyword.endswith('ಅದು') or \
        #             keyword.endswith('ವು') or keyword.endswith('ನ') or keyword.endswith('ನಿಂದ') or \
        #             keyword.endswith('ರು') or keyword.endswith('ದು') or keyword.endswith('ಕೆ') or \
        #             keyword.endswith('ಬೇಕು'):
        #         keyword_type = "noun"
        #
        #     # Assign the keyword to the appropriate category
        #     if keyword_type == "noun":
        #         noun_words.append(keyword)
        #
        #     # Check if the keyword is a verb or pronoun based on specific conditions
        #     if keyword.endswith('ನು') or keyword.endswith('ನಿಂದ') or keyword.endswith('ಗಳು') or \
        #             keyword.endswith('ತೆ') or keyword.endswith('ದು') or keyword.endswith('ಕೆ') or \
        #             keyword.endswith('ಬೇಕು'):
        #         keyword_type = "verb"
        #
        #     elif keyword.endswith('ಅವರು') or keyword.endswith('ಅವಳು') or keyword.endswith('ಅದು') or \
        #             keyword.endswith('ನು') or keyword.endswith('ನಿಂದ') or keyword.endswith('ಗಳು') or \
        #             keyword.endswith('ಇವನು') or keyword.endswith('ರು'):
        #         keyword_type = "pronoun"


        for keyword in keywords:
            keyword_type = None

            # Check if the keyword is a noun based on specific conditions
            if any(keyword.endswith(suffix) for suffix in
                   ['ವರು', 'ವಳು', 'ಅದು', 'ವು', 'ನ', 'ನಿಂದ', 'ರು', 'ದು', 'ಕೆ', 'ಬೇಕು', 'ನಿಂದ', 'ಆಯಿ', 'ಉತ್ತಿತ್ತು',
                    'ಇದೆ', 'ಇಲ್ಲ', 'ಇರುತ್ತದೆ', 'ಒಂದು', 'ಒನ್ನು', 'ಗಳಿಗೆ', 'ಚೆನ್ನಾಗಿ', 'ತ್ತದೆ', 'ತ್ತಾನೆ', 'ದೆಯ', 'ದೇವರೆ','ಬೇಕು',
                    'ನಿಂದ', 'ಇಸುವುದು', 'ಇಕ್ಕುವುದು', 'ಇಕ್ಕೊಳ್ಳುವುದು', 'ಆಗುವುದು', 'ಕೊಳ್ಳುವುದು', 'ಮಾಡುವುದು', 'ಹೋಗುವುದು',
                    'ಬರುವುದು', 'ಹೇಳಿದು', 'ತಿಳಿಯುವುದು', 'ನೋಡುವುದು', 'ಕೇಳುವುದು', 'ಕೊಡುವುದು', 'ಬರುಸು', 'ಹೇಳಿಸೆ',
                    'ತಿಳಿದ್ದೆ', 'ನೋಡಿದ್ದೆ', 'ಕೇಳ್ದೆ', ' ಕೊಟ್ಟೆ', ' ಅಲ್ಲಿ', 'ಇಲ್ಲಿ', 'ಇದಿಗೆ, ಅವರು,ನಮ್ಮ',' ತನ್ನ', 'ಇವನು', 'ಅವಳು',
                    'ಅದು', 'ಇವರು', 'ಅಂತ', 'ಎಲ್ಲಾ', ' ಎಂತ', ' ಎಲ್ಲವೋ', 'ಎಂತದ್ದೇ', 'ಎಲ್ಲಾದು']):
                keyword_type = "noun"
            # Check if the keyword is a noun based on specific prefix conditions
            if any(keyword.startswith(prefix) for prefix in
                   ['ಶಿಕ್ಷಕ', 'ಹೆಣ್ಣು', 'ಗಾಳಿ', 'ಹಣ', 'ಸ್ಥಳ', 'ಗಣಿತ', 'ಹಾಲು', 'ಹೊಟ್ಟೆ',"ಅತಿ", "ಅನು", "ಅಂತರ", "ಏಕ", "ಕಿರು",
                    "ಕುಲ", "ಕೃತ", "ಗುಪ್ತ", "ಚತುರ್", "ಜನ", "ತೃತೀಯ", "ದಶ", "ಧರ್ಮ", "ನವ", "ಪಂಚ", "ಪರ", "ಪುನರ್", "ಬಹು",
                    "ಸಹ", "ಸಮ", "ಸಾರ್ವ"]):
                keyword_type = "noun"


            # Check if the keyword is a verb based on specific conditions
            elif any(keyword.endswith(suffix) for suffix in ['ನು', 'ನಿಂದ', 'ತೆ', 'ದು', 'ಕೆ', 'ಬೇಕು']):
                keyword_type = "verb"

            # Check if the keyword is a pronoun based on specific conditions
            elif any(keyword.endswith(suffix) for suffix in ['ಅವರು', 'ಅವಳು', 'ಅದು', 'ನು', 'ನಿಂದ', 'ಗಳು', 'ಇವನು', 'ರು']):
                keyword_type = "pronoun"

            # Additional checks based on prefixes
            if any(keyword.startswith(prefix) for prefix in ['ಹೇಳು', 'ನೋಡು', 'ಹೋಗು', 'ಬರು']):
                keyword_type = "verb"

            # Assign the keyword to the appropriate category
            if keyword_type == "noun":
                noun_words.append(keyword)
            elif keyword_type == "verb":
                verb_words.append(keyword)
            elif keyword_type == "pronoun":
                pronoun_words.append(keyword)

        return noun_words, verb_words, pronoun_words


if __name__ == '__main__':
    app = wx.App()
    KannadaKeywordExtractorApp(None, title='Kannada Keyword Extractor')
    app.MainLoop()
